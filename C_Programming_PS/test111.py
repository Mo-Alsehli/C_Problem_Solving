from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Set the font to a commonly used Arabic font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Add a custom heading style for Arabic
heading_style = doc.styles.add_style('ArabicHeading', 1)
heading_font = heading_style.font
heading_font.name = 'Arial'
heading_font.size = Pt(14)
heading_font.bold = True

# Function to set Arabic alignment
def set_arabic_alignment(paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = 2  # Right alignment for Arabic text

# Add Title
title = doc.add_heading('مشروع صالة الألعاب الرياضية', 0)
set_arabic_alignment(title)

# Add Project Information Table
project_info = doc.add_heading('معلومات المشروع', level=1)
project_info.style = doc.styles['ArabicHeading']
set_arabic_alignment(project_info)

table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'رقم المشروع'
hdr_cells[1].text = 'الاسم'
hdr_cells[2].text = 'الرقم'

data = [
    ("203205", "عبدالرحمن اسامة محمد عبدالغفار", ""),
    ("203203", "احمد مسعد عراقيب", ""),
    ("203226", "محمد مجدي محمد اليماني السحلي", ""),
]

for project_number, name, number in data:
    row_cells = table.add_row().cells
    row_cells[0].text = project_number
    row_cells[1].text = name
    row_cells[2].text = number

# Set alignment for table cells
for row in table.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.alignment = 2

# Add Project Costs
project_costs = doc.add_heading('التكلفة المتوقعة للمشروع:', level=1)
project_costs.style = doc.styles['ArabicHeading']
set_arabic_alignment(project_costs)

paragraphs = [
    "1. تكلفة الايجار للمكان: 800$ شهريا",
    "2. تكلفة الاجهزة: 50000$",
    "3. تكلفة الصيانة سنويا: 2000$",
    "4. تكاليف المدرب: 30$ يوميا",
    "5. تكلفة الكهرباء: 5$ يوميا",
    "التكلفة الكلية في البداية: 100000$"
]

for text in paragraphs:
    p = doc.add_paragraph(text)
    set_arabic_alignment(p)

# Add Loan and Repayment
loan_repayment = doc.add_heading('القرض والتسديد', level=1)
loan_repayment.style = doc.styles['ArabicHeading']
set_arabic_alignment(loan_repayment)

loan_text = "تم الاتفاق على أخذ قرض بقيمة 100000$ من البنك على أن يتم سداد القرض على خمس سنوات ابتداء من يناير 2023 حتى يناير 2028، بفائدة 10% سنويا."
loan_paragraph = doc.add_paragraph(loan_text)
set_arabic_alignment(loan_paragraph)

# Add Repayment Details Table
repayment_details = doc.add_heading('تفاصيل التسديد:', level=2)
repayment_details.style = doc.styles['ArabicHeading']
set_arabic_alignment(repayment_details)

repayment_table = doc.add_table(rows=1, cols=4)
hdr_cells = repayment_table.rows[0].cells
hdr_cells[0].text = 'السنة'
hdr_cells[1].text = 'الرصيد'
hdr_cells[2].text = 'الفائدة'
hdr_cells[3].text = 'المبلغ المستحق'

repayment_data = [
    ("1", "100000", "10000", "10000"),
    ("2", "100000", "10000", "10000"),
    ("3", "100000", "10000", "10000"),
    ("4", "100000", "10000", "10000"),
    ("5", "100000", "10000", "110000"),
]

for year, balance, interest, due_amount in repayment_data:
    row_cells = repayment_table.add_row().cells
    row_cells[0].text = year
    row_cells[1].text = balance
    row_cells[2].text = interest
    row_cells[3].text = due_amount

# Set alignment for table cells
for row in repayment_table.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.alignment = 2

total_paragraph = doc.add_paragraph("الإجمالي: 150000$")
set_arabic_alignment(total_paragraph)

# Add Total Project Costs Calculation
total_costs = doc.add_heading('حساب التكاليف الكلية للمشروع', level=1)
total_costs.style = doc.styles['ArabicHeading']
set_arabic_alignment(total_costs)

fixed_costs = doc.add_heading('1. التكاليف الثابتة:', level=2)
fixed_costs.style = doc.styles['ArabicHeading']
set_arabic_alignment(fixed_costs)

fixed_paragraphs = [
    "- إيجار المكان: 800$ شهريا",
    "- قسط شراء الأجهزة (ثمنها 50000$ على خمس سنوات بفائدة 8%): A = P * A/P = 50000 * 0.2505 = 12525$",
    "- تكلفة الصيانة: 2000$ سنويا"
]

for text in fixed_paragraphs:
    p = doc.add_paragraph(text)
    set_arabic_alignment(p)

variable_costs = doc.add_heading('2. التكاليف المتغيرة:', level=2)
variable_costs.style = doc.styles['ArabicHeading']
set_arabic_alignment(variable_costs)

variable_paragraphs = [
    "- الكهرباء: 5$ يوميا",
    "- العمالة: 30$ يوميا",
    "Tc = Cf + Cv = (800 * 12 + 12525 + 2000) + (35 * 365) = 36900$",
    "الإيرادات المتوقعة:",
    "- 110$ يوميا",
    "الإجمالي السنوي: 40150$"
]

for text in variable_paragraphs:
    p = doc.add_paragraph(text)
    set_arabic_alignment(p)

# Add Break-Even Calculation
break_even = doc.add_heading('حساب نقطة التعادل:', level=1)
break_even.style = doc.styles['ArabicHeading']
set_arabic_alignment(break_even)

break_even_paragraphs = [
    "I = Tc",
    "I = Cf + Cv",
    "V * R = Cf + V * N",
    "N = 1.388 سنة أي تقريبا سنة وخمس شهور"
]

for text in break_even_paragraphs:
    p = doc.add_paragraph(text)
    set_arabic_alignment(p)

# Add Expected Profit Calculation
expected_profit_5_years = doc.add_heading('الربح المتوقع بعد خمس سنوات:', level=1)
expected_profit_5_years.style = doc.styles['ArabicHeading']
set_arabic_alignment(expected_profit_5_years)

profit_5_years_text = "P = I - Tc = 40150 * 5 - (22415.77 + 12775 * 5) = 114459.23$"
profit_5_years_paragraph = doc.add_paragraph(profit_5_years_text)
set_arabic_alignment(profit_5_years_paragraph)

expected_profit_10_years = doc.add_heading('الربح المتوقع بعد عشر سنوات:', level=1)
expected_profit_10_years.style = doc.styles['ArabicHeading']
set_arabic_alignment(expected_profit_10_years)

profit_10_years_text = "P = I - Tc = 40150 * 10 - (22415.77 + 12775 * 10) = 251334.23$"
profit_10_years_paragraph_1 = doc.add_paragraph(profit_10_years_text)
set_arabic_alignment(profit_10_years_paragraph_1)

profit_10_years_paragraph_2 = doc.add_paragraph("بعد خصم قيمة القرض:")
set_arabic_alignment(profit_10_years_paragraph_2)

profit_10_years_paragraph_3 = doc.add_paragraph("251334.23 - 150000 = 101334.23$")
set_arabic_alignment(profit_10_years_paragraph_3)

# Add Depreciation and Equipment Renewal
depreciation = doc.add_heading('معدل الاستهلاك وتجديد الأجهزة:', level=1)
depreciation.style = doc.styles['ArabicHeading']
set_arabic_alignment(depreciation)

depreciation_paragraphs = [
    "من المتوقع حدوث تجديد وإحلال في الماكينات والأجهزة بعد 15 سنة وسعر بيعها مستهلكة 20000$",
    "D = (I - S) / N = (50000 - 20000) / 15 = 2000$ / سنة",
    "القيمة الافتراضية في السنة العاشرة:",
    "Bvi = I - Icd = 50000 - 10 * 2000 = 30000$",
    "معدل الاستهلاك السنوي:",
    "R = 100 / 15 = 6.66%"
]

for text in depreciation_paragraphs:
    p = doc.add_paragraph(text)
    set_arabic_alignment(p)

# Save Document
doc.save("C:/Users/dell/Downloads/Organized_Project_Document1.docx")
